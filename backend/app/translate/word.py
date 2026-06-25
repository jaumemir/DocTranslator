# translate/word.py
"""
Chunking strategy:
- Body text is processed by paragraph
- Tables are processed by cell
- Headers and footers are processed separately
- Non-text elements like images and charts are preserved
- Paragraph formatting including alignment, indentation, and line spacing is maintained
"""
import datetime
import logging
import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from threading import Event
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
from . import to_translate
from . import common

# Chunking configuration
MAX_CHUNK_SIZE = 2000


@dataclass
class RunStyle:
    """Run style data"""
    font_name: Optional[str] = None
    font_name_east_asia: Optional[str] = None
    font_size: Optional[Pt] = None
    font_bold: Optional[bool] = None
    font_italic: Optional[bool] = None
    font_underline: Optional[bool] = None
    font_color_rgb: Optional[RGBColor] = None
    font_strike: Optional[bool] = None


@dataclass
class ParagraphFormat:
    """Paragraph format data"""
    alignment: Optional[int] = None
    left_indent: Optional[int] = None
    right_indent: Optional[int] = None
    first_line_indent: Optional[int] = None
    space_before: Optional[int] = None
    space_after: Optional[int] = None
    line_spacing: Optional[float] = None
    line_spacing_rule: Optional[int] = None


@dataclass
class TextBlock:
    """Text block"""
    uid: str
    block_type: str  # paragraph, table_cell, header, footer

    # Position
    paragraph_index: int = -1
    table_index: int = -1
    row_index: int = -1
    col_index: int = -1
    header_footer_type: str = ""
    section_index: int = -1

    # Content
    original_text: str = ""
    translated_text: str = ""

    # Style
    run_style: Optional[RunStyle] = None
    para_format: Optional[ParagraphFormat] = None

    # Status
    complete: bool = False
    skip: bool = False
    count: int = 0

    # Sub-block information
    is_sub: bool = False
    sub_index: int = 0
    sub_total: int = 1
    parent_uid: str = ""


def start(trans: Dict[str, Any]) -> bool:
    """Word document translation entry point"""
    translate_id = trans['id']
    start_time = datetime.datetime.now()

    trans_type = trans.get('type', 'trans_only_inherit')
    only_translation = 'only' in trans_type
    inherit_format = 'inherit' in trans_type

    logging.info(
        f"[Task {translate_id}] Translation mode: only={only_translation}, inherit={inherit_format}")

    try:
        document = Document(trans['file_path'])
    except Exception as e:
        logging.error(f"[Task {translate_id}] Unable to open document: {e}")
        to_translate.error(translate_id, f"Unable to open document: {str(e)}")
        return False

    try:
        text_blocks = _extract_all_text_blocks(document)
    except Exception as e:
        logging.error(f"[Task {translate_id}] Text extraction failed: {e}")
        to_translate.error(translate_id, f"Text extraction failed: {str(e)}")
        return False

    blocks_to_translate = [b for b in text_blocks if not b.skip]

    if not blocks_to_translate:
        logging.info(f"[Task {translate_id}] No text to translate in document")
        document.save(trans['target_file'])
        to_translate.complete(trans, 0, "0s")
        return True

    logging.info(
        f"[Task {translate_id}] Total {len(text_blocks)} blocks, {len(blocks_to_translate)} need translation")

    texts = _blocks_to_texts(blocks_to_translate)

    event = Event()
    success = to_translate.translate_batch(trans, texts, event)
    if not success:
        return False

    _sync_results(blocks_to_translate, texts)

    try:
        text_count = _apply_translation(document, text_blocks, only_translation,
                                        inherit_format, trans.get('lang', 'English'))
        document.save(trans['target_file'])
    except Exception as e:
        logging.error(f"[Task {translate_id}] Failed to save document: {e}")
        to_translate.error(translate_id, f"Failed to save document: {str(e)}")
        return False

    end_time = datetime.datetime.now()
    spend_time = common.display_spend(start_time, end_time)
    to_translate.complete(trans, text_count, spend_time)
    return True


# ==================== Text Extraction ====================

def _extract_all_text_blocks(document: Document) -> List[TextBlock]:
    """Extract all text blocks from document"""
    blocks = []
    uid_counter = [0]

    def next_uid(prefix: str) -> str:
        uid_counter[0] += 1
        return f"{prefix}_{uid_counter[0]}"

    for para_idx, paragraph in enumerate(document.paragraphs):
        para_blocks = _extract_paragraph_blocks(paragraph, para_idx, next_uid)
        blocks.extend(para_blocks)

    for table_idx, table in enumerate(document.tables):
        table_blocks = _extract_table_blocks(table, table_idx, next_uid)
        blocks.extend(table_blocks)

    for section_idx, section in enumerate(document.sections):
        hf_blocks = _extract_header_footer_blocks(section, section_idx, next_uid)
        blocks.extend(hf_blocks)

    return blocks


def _extract_paragraph_blocks(paragraph: Paragraph, para_idx: int, next_uid) -> List[TextBlock]:
    """Extract paragraph text blocks"""
    blocks = []

    text = _get_paragraph_text(paragraph)

    if not text or not text.strip():
        return blocks

    # Extract paragraph format
    para_format = _extract_paragraph_format(paragraph)

    if not _should_translate(text):
        block = TextBlock(
            uid=next_uid("para"),
            block_type="paragraph",
            paragraph_index=para_idx,
            original_text=text,
            para_format=para_format,
            skip=True
        )
        blocks.append(block)
        return blocks

    run_style = _extract_first_run_style(paragraph)

    if len(text) <= MAX_CHUNK_SIZE:
        block = TextBlock(
            uid=next_uid("para"),
            block_type="paragraph",
            paragraph_index=para_idx,
            original_text=text,
            run_style=run_style,
            para_format=para_format
        )
        blocks.append(block)
    else:
        sub_texts = _split_by_sentences(text, MAX_CHUNK_SIZE)
        parent_uid = next_uid("para_parent")
        for i, sub_text in enumerate(sub_texts):
            block = TextBlock(
                uid=next_uid("para_sub"),
                block_type="paragraph",
                paragraph_index=para_idx,
                original_text=sub_text,
                run_style=run_style,
                para_format=para_format,
                is_sub=True,
                sub_index=i,
                sub_total=len(sub_texts),
                parent_uid=parent_uid
            )
            blocks.append(block)

    return blocks


def _extract_table_blocks(table: Table, table_idx: int, next_uid) -> List[TextBlock]:
    """Extract table text blocks"""
    blocks = []
    processed_cells = set()

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_id = id(cell._tc)
            if cell_id in processed_cells:
                continue
            processed_cells.add(cell_id)

            text = _get_cell_text(cell)

            if not text or not text.strip():
                continue

            # Extract format from first paragraph of cell
            para_format = None
            if cell.paragraphs:
                para_format = _extract_paragraph_format(cell.paragraphs[0])

            if not _should_translate(text):
                block = TextBlock(
                    uid=next_uid("cell"),
                    block_type="table_cell",
                    table_index=table_idx,
                    row_index=row_idx,
                    col_index=col_idx,
                    original_text=text,
                    para_format=para_format,
                    skip=True
                )
                blocks.append(block)
                continue

            run_style = _extract_cell_first_run_style(cell)

            if len(text) <= MAX_CHUNK_SIZE:
                block = TextBlock(
                    uid=next_uid("cell"),
                    block_type="table_cell",
                    table_index=table_idx,
                    row_index=row_idx,
                    col_index=col_idx,
                    original_text=text,
                    run_style=run_style,
                    para_format=para_format
                )
                blocks.append(block)
            else:
                sub_texts = _split_by_sentences(text, MAX_CHUNK_SIZE)
                parent_uid = next_uid("cell_parent")
                for i, sub_text in enumerate(sub_texts):
                    block = TextBlock(
                        uid=next_uid("cell_sub"),
                        block_type="table_cell",
                        table_index=table_idx,
                        row_index=row_idx,
                        col_index=col_idx,
                        original_text=sub_text,
                        run_style=run_style,
                        para_format=para_format,
                        is_sub=True,
                        sub_index=i,
                        sub_total=len(sub_texts),
                        parent_uid=parent_uid
                    )
                    blocks.append(block)

    return blocks


def _extract_header_footer_blocks(section, section_idx: int, next_uid) -> List[TextBlock]:
    """Extract header and footer text blocks"""
    blocks = []

    hf_types = [
        ('header', section.header),
        ('footer', section.footer),
        ('first_header', section.first_page_header),
        ('first_footer', section.first_page_footer),
    ]

    for hf_type, hf in hf_types:
        if hf is None:
            continue
        try:
            if hf.is_linked_to_previous:
                continue
        except:
            pass

        try:
            for para_idx, paragraph in enumerate(hf.paragraphs):
                text = _get_paragraph_text(paragraph)
                if text and text.strip() and _should_translate(text):
                    run_style = _extract_first_run_style(paragraph)
                    para_format = _extract_paragraph_format(paragraph)
                    block = TextBlock(
                        uid=next_uid(f"hf_{hf_type}"),
                        block_type="header_footer",
                        header_footer_type=hf_type,
                        section_index=section_idx,
                        paragraph_index=para_idx,
                        original_text=text,
                        run_style=run_style,
                        para_format=para_format
                    )
                    blocks.append(block)
        except Exception as e:
            logging.warning(f"Failed to extract header/footer: {e}")

    return blocks


def _extract_paragraph_format(paragraph: Paragraph) -> ParagraphFormat:
    """Extract paragraph format"""
    fmt = ParagraphFormat()

    try:
        pf = paragraph.paragraph_format

        # Alignment
        fmt.alignment = paragraph.alignment

        # Indentation
        fmt.left_indent = pf.left_indent
        fmt.right_indent = pf.right_indent
        fmt.first_line_indent = pf.first_line_indent

        # Space before/after
        fmt.space_before = pf.space_before
        fmt.space_after = pf.space_after

        # Line spacing
        fmt.line_spacing = pf.line_spacing
        fmt.line_spacing_rule = pf.line_spacing_rule

    except Exception as e:
        logging.warning(f"Failed to extract paragraph format: {e}")

    return fmt


def _get_paragraph_text(paragraph: Paragraph) -> str:
    """Get paragraph text"""
    texts = []
    for run in paragraph.runs:
        if _run_has_image(run):
            continue
        if run.text:
            texts.append(run.text)
    return ''.join(texts)


def _get_cell_text(cell: _Cell) -> str:
    """Get cell text"""
    texts = []
    for paragraph in cell.paragraphs:
        para_text = _get_paragraph_text(paragraph)
        if para_text:
            texts.append(para_text)
    return '\n'.join(texts)


def _run_has_image(run: Run) -> bool:
    """Check if run contains an image"""
    try:
        drawings = run._element.findall('.//' + qn('w:drawing'))
        if drawings:
            return True
        picts = run._element.findall('.//' + qn('w:pict'))
        if picts:
            return True
    except:
        pass
    return False


def _extract_first_run_style(paragraph: Paragraph) -> Optional[RunStyle]:
    """Extract style from first valid run in paragraph"""
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return _extract_run_style(run)
    return None


def _extract_cell_first_run_style(cell: _Cell) -> Optional[RunStyle]:
    """Extract style from first valid run in cell"""
    for paragraph in cell.paragraphs:
        style = _extract_first_run_style(paragraph)
        if style:
            return style
    return None


def _extract_run_style(run: Run) -> RunStyle:
    """Extract run style"""
    style = RunStyle()

    try:
        font = run.font
        style.font_name = font.name
        style.font_size = font.size
        style.font_bold = font.bold
        style.font_italic = font.italic
        style.font_underline = font.underline
        style.font_strike = font.strike

        try:
            rPr = run._element.rPr
            if rPr is not None:
                rFonts = rPr.rFonts
                if rFonts is not None:
                    style.font_name_east_asia = rFonts.get(qn('w:eastAsia'))
        except:
            pass

        try:
            if font.color and font.color.rgb:
                style.font_color_rgb = font.color.rgb
        except:
            pass

    except Exception as e:
        logging.warning(f"Failed to extract run style: {e}")

    return style


def _should_translate(text: str) -> bool:
    """Determine if text should be translated"""
    if not text or not text.strip():
        return False

    text = text.strip()

    if common.is_all_punc(text):
        return False

    if re.match(r'^[\d\s\.\-\+\*\/\=\%\(\)\[\]\{\}]+$', text):
        return False

    if re.match(r'^(第\s*\d+\s*页|Page\s+\d+|\d+\s*/\s*\d+)$', text, re.IGNORECASE):
        return False

    if len(text) < 2:
        return False

    return True


def _split_by_sentences(text: str, max_size: int) -> List[str]:
    """Split long text by sentence boundaries"""
    sentence_endings = r'([.!?。！？；;]\s*)'

    parts = re.split(sentence_endings, text)

    sentences = []
    i = 0
    while i < len(parts):
        sentence = parts[i]
        if i + 1 < len(parts) and re.match(sentence_endings, parts[i + 1]):
            sentence += parts[i + 1]
            i += 2
        else:
            i += 1
        if sentence.strip():
            sentences.append(sentence)

    chunks = []
    current = ""

    for sentence in sentences:
        if len(sentence) > max_size:
            if current:
                chunks.append(current)
                current = ""
            for j in range(0, len(sentence), max_size):
                chunks.append(sentence[j:j + max_size])
            continue

        if len(current) + len(sentence) <= max_size:
            current += sentence
        else:
            if current:
                chunks.append(current)
            current = sentence

    if current:
        chunks.append(current)

    return chunks if chunks else [text]


# ==================== Translation Interface ====================

def _blocks_to_texts(blocks: List[TextBlock]) -> List[Dict]:
    """Convert TextBlock list to translation interface format"""
    return [{
        'text': b.original_text,
        'original': b.original_text,
        'complete': False,
        'count': 0,
        '_block_uid': b.uid
    } for b in blocks]


def _sync_results(blocks: List[TextBlock], texts: List[Dict]):
    """Sync translation results to TextBlock"""
    block_map = {b.uid: b for b in blocks}

    for text_item in texts:
        uid = text_item.get('_block_uid')
        if uid and uid in block_map:
            block = block_map[uid]
            block.translated_text = text_item.get('text', block.original_text)
            block.complete = text_item.get('complete', False)
            block.count = text_item.get('count', 0)


# ==================== Apply Translation ====================

def _apply_translation(document: Document, all_blocks: List[TextBlock],
                       only_translation: bool, inherit_format: bool,
                       target_lang: str) -> int:
    """Apply translation results to document"""
    text_count = 0

    para_blocks = _organize_paragraph_blocks(all_blocks)
    table_blocks = _organize_table_blocks(all_blocks)
    hf_blocks = _organize_header_footer_blocks(all_blocks)

    for para_idx, paragraph in enumerate(document.paragraphs):
        if para_idx in para_blocks:
            blocks = para_blocks[para_idx]
            count = _apply_to_paragraph(paragraph, blocks, only_translation,
                                        inherit_format, target_lang)
            text_count += count

    for table_idx, table in enumerate(document.tables):
        processed_cells = set()
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id in processed_cells:
                    continue
                processed_cells.add(cell_id)

                key = (table_idx, row_idx, col_idx)
                if key in table_blocks:
                    blocks = table_blocks[key]
                    count = _apply_to_cell(cell, blocks, only_translation,
                                           inherit_format, target_lang)
                    text_count += count

    for section_idx, section in enumerate(document.sections):
        _apply_to_header_footer(section, section_idx, hf_blocks,
                                only_translation, inherit_format, target_lang)

    return text_count


def _organize_paragraph_blocks(blocks: List[TextBlock]) -> Dict[int, List[TextBlock]]:
    """Organize blocks by paragraph index"""
    result = {}
    for b in blocks:
        if b.block_type == "paragraph":
            if b.paragraph_index not in result:
                result[b.paragraph_index] = []
            result[b.paragraph_index].append(b)

    for idx in result:
        result[idx].sort(key=lambda x: x.sub_index)

    return result


def _organize_table_blocks(blocks: List[TextBlock]) -> Dict[tuple, List[TextBlock]]:
    """Organize blocks by table position"""
    result = {}
    for b in blocks:
        if b.block_type == "table_cell":
            key = (b.table_index, b.row_index, b.col_index)
            if key not in result:
                result[key] = []
            result[key].append(b)

    for key in result:
        result[key].sort(key=lambda x: x.sub_index)

    return result


def _organize_header_footer_blocks(blocks: List[TextBlock]) -> Dict[tuple, List[TextBlock]]:
    """Organize blocks by header/footer"""
    result = {}
    for b in blocks:
        if b.block_type == "header_footer":
            key = (b.section_index, b.header_footer_type, b.paragraph_index)
            if key not in result:
                result[key] = []
            result[key].append(b)

    return result


def _apply_to_paragraph(paragraph: Paragraph, blocks: List[TextBlock],
                        only_translation: bool, inherit_format: bool,
                        target_lang: str) -> int:
    """Apply translation to paragraph"""
    text_count = sum(b.count for b in blocks)

    # Merge sub-blocks
    if any(b.is_sub for b in blocks):
        original = ''.join(b.original_text for b in blocks)
        translated = ''.join(b.translated_text or b.original_text for b in blocks)
    else:
        original = blocks[0].original_text
        translated = blocks[0].translated_text or original

    if blocks[0].skip:
        return 0

    run_style = blocks[0].run_style
    para_format = blocks[0].para_format

    if only_translation:
        _replace_paragraph_text(paragraph, translated, run_style, para_format,
                                inherit_format, target_lang, len(original))
    else:
        _append_translation_to_paragraph(paragraph, translated, run_style, para_format,
                                         inherit_format, target_lang)

    return text_count


def _apply_to_cell(cell: _Cell, blocks: List[TextBlock],
                   only_translation: bool, inherit_format: bool,
                   target_lang: str) -> int:
    """Apply translation to cell"""
    text_count = sum(b.count for b in blocks)

    if any(b.is_sub for b in blocks):
        original = ''.join(b.original_text for b in blocks)
        translated = ''.join(b.translated_text or b.original_text for b in blocks)
    else:
        original = blocks[0].original_text
        translated = blocks[0].translated_text or original

    if blocks[0].skip:
        return 0

    run_style = blocks[0].run_style
    para_format = blocks[0].para_format

    if only_translation:
        _replace_cell_text(cell, translated, run_style, para_format,
                           inherit_format, target_lang, len(original))
    else:
        _append_translation_to_cell(cell, translated, run_style, para_format,
                                    inherit_format, target_lang)

    return text_count


def _apply_to_header_footer(section, section_idx: int,
                            hf_blocks: Dict[tuple, List[TextBlock]],
                            only_translation: bool, inherit_format: bool,
                            target_lang: str):
    """Apply translation to header/footer"""
    hf_map = {
        'header': section.header,
        'footer': section.footer,
        'first_header': section.first_page_header,
        'first_footer': section.first_page_footer,
    }

    for hf_type, hf in hf_map.items():
        if hf is None:
            continue
        try:
            for para_idx, paragraph in enumerate(hf.paragraphs):
                key = (section_idx, hf_type, para_idx)
                if key in hf_blocks:
                    blocks = hf_blocks[key]
                    translated = blocks[0].translated_text or blocks[0].original_text
                    original_len = len(blocks[0].original_text)

                    if only_translation:
                        _replace_paragraph_text(paragraph, translated,
                                                blocks[0].run_style, blocks[0].para_format,
                                                inherit_format, target_lang, original_len)
                    else:
                        _append_translation_to_paragraph(paragraph, translated,
                                                         blocks[0].run_style, blocks[0].para_format,
                                                         inherit_format, target_lang)
        except Exception as e:
            logging.warning(f"Failed to process header/footer: {e}")


def _replace_paragraph_text(paragraph: Paragraph, new_text: str,
                            run_style: Optional[RunStyle],
                            para_format: Optional[ParagraphFormat],
                            inherit_format: bool, target_lang: str,
                            original_len: int):
    """Replace paragraph text (translation-only mode)"""

    # Save image elements
    image_elements = []
    for run in paragraph.runs:
        if _run_has_image(run):
            image_elements.append(deepcopy(run._element))

    if inherit_format and paragraph.runs:
        # Inherit format: only replace text content
        for run in paragraph.runs:
            if not _run_has_image(run):
                run.text = ""

        # Set translation in first non-image run
        first_text_run = None
        for run in paragraph.runs:
            if not _run_has_image(run):
                first_text_run = run
                break

        if first_text_run:
            first_text_run.text = new_text
            _adjust_font_size_for_run(first_text_run, run_style, len(new_text), original_len)
        else:
            # No text run, create new one
            new_run = paragraph.add_run(new_text)
            _apply_run_style(new_run, run_style, target_lang, len(new_text), original_len)
    else:
        # Reformat: clear and rebuild
        paragraph.clear()

        # Add translation
        new_run = paragraph.add_run(new_text)
        _apply_run_style(new_run, run_style, target_lang, len(new_text), original_len)

        # Restore images
        for img_elem in image_elements:
            paragraph._p.append(img_elem)

    # Restore paragraph format
    _apply_paragraph_format(paragraph, para_format)


def _append_translation_to_paragraph(paragraph: Paragraph, translated: str,
                                     run_style: Optional[RunStyle],
                                     para_format: Optional[ParagraphFormat],
                                     inherit_format: bool, target_lang: str):
    """Append translation to end of paragraph (original+translation mode)"""

    # Find last non-image run
    last_text_run = None
    for run in reversed(paragraph.runs):
        if not _run_has_image(run):
            last_text_run = run
            break

    if last_text_run is None:
        if paragraph.runs:
            last_text_run = paragraph.runs[-1]
        else:
            # Empty paragraph, add directly
            new_run = paragraph.add_run(translated)
            _apply_run_style(new_run, run_style, target_lang, len(translated), 0)
            _apply_paragraph_format(paragraph, para_format)
            return

    # Add line break
    last_text_run.add_break()

    # Add translation run
    new_run = paragraph.add_run(translated)

    if inherit_format and run_style:
        _apply_run_style(new_run, run_style, target_lang, len(translated), 0)
    else:
        _apply_run_style(new_run, run_style, target_lang, len(translated), 0)

    # Ensure paragraph format is maintained
    _apply_paragraph_format(paragraph, para_format)


def _replace_cell_text(cell: _Cell, new_text: str,
                       run_style: Optional[RunStyle],
                       para_format: Optional[ParagraphFormat],
                       inherit_format: bool, target_lang: str,
                       original_len: int):
    """Replace cell text (translation-only mode)"""

    # Save images
    image_elements = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if _run_has_image(run):
                image_elements.append(deepcopy(run._element))

    # Clear cell content but keep first paragraph
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    # Clear content of first paragraph
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run.clear()

    # Process multi-line text
    lines = new_text.split('\n')

    for i, line in enumerate(lines):
        if i == 0:
            para = first_para
            # Clear existing runs
            for run in para.runs:
                run.text = ""
            # Add new text
            if para.runs:
                para.runs[0].text = line
                _adjust_font_size_for_run(para.runs[0], run_style, len(line), original_len)
            else:
                new_run = para.add_run(line)
                _apply_run_style(new_run, run_style, target_lang, len(line), original_len)
        else:
            para = cell.add_paragraph()
            new_run = para.add_run(line)
            _apply_run_style(new_run, run_style, target_lang, len(line), 0)

        # Apply paragraph format
        _apply_paragraph_format(para, para_format)

    # Restore images to last paragraph
    if image_elements and cell.paragraphs:
        for img_elem in image_elements:
            cell.paragraphs[-1]._p.append(img_elem)


def _append_translation_to_cell(cell: _Cell, translated: str,
                                run_style: Optional[RunStyle],
                                para_format: Optional[ParagraphFormat],
                                inherit_format: bool, target_lang: str):
    """Append translation to end of cell (original+translation mode)"""

    if not cell.paragraphs:
        para = cell.add_paragraph()
        new_run = para.add_run(translated)
        _apply_run_style(new_run, run_style, target_lang, len(translated), 0)
        _apply_paragraph_format(para, para_format)
        return

    last_para = cell.paragraphs[-1]

    # Add line break
    if last_para.runs:
        last_para.runs[-1].add_break()

    # Add translation
    new_run = last_para.add_run(translated)
    _apply_run_style(new_run, run_style, target_lang, len(translated), 0)

    # Ensure format is maintained
    _apply_paragraph_format(last_para, para_format)


def _apply_paragraph_format(paragraph: Paragraph, fmt: Optional[ParagraphFormat]):
    """Apply paragraph format"""
    if not fmt:
        return

    try:
        # Alignment
        if fmt.alignment is not None:
            paragraph.alignment = fmt.alignment

        pf = paragraph.paragraph_format

        # Indentation
        if fmt.left_indent is not None:
            pf.left_indent = fmt.left_indent
        if fmt.right_indent is not None:
            pf.right_indent = fmt.right_indent
        if fmt.first_line_indent is not None:
            pf.first_line_indent = fmt.first_line_indent

        # Space before/after
        if fmt.space_before is not None:
            pf.space_before = fmt.space_before
        if fmt.space_after is not None:
            pf.space_after = fmt.space_after

        # Line spacing
        if fmt.line_spacing is not None:
            pf.line_spacing = fmt.line_spacing
        if fmt.line_spacing_rule is not None:
            pf.line_spacing_rule = fmt.line_spacing_rule

    except Exception as e:
        logging.warning(f"Failed to apply paragraph format: {e}")


def _apply_run_style(run: Run, style: Optional[RunStyle],
                     target_lang: str, new_len: int, original_len: int):
    """Apply style to run"""
    if not style:
        _set_default_font(run, target_lang)
        return

    try:
        font = run.font

        # Font name
        if style.font_name:
            font.name = style.font_name

        # East Asian font
        if target_lang in ['中文', '日语', '韩语']:
            try:
                run._element.get_or_add_rPr()
                rFonts = run._element.rPr.get_or_add_rFonts()
                if style.font_name_east_asia:
                    rFonts.set(qn('w:eastAsia'), style.font_name_east_asia)
                else:
                    rFonts.set(qn('w:eastAsia'), '微软雅黑')
            except:
                pass

        # Font size
        if style.font_size:
            font.size = _calculate_adjusted_size(style.font_size, new_len, original_len)

        # Bold, italic, underline, strikethrough
        if style.font_bold is not None:
            font.bold = style.font_bold
        if style.font_italic is not None:
            font.italic = style.font_italic
        if style.font_underline is not None:
            font.underline = style.font_underline
        if style.font_strike is not None:
            font.strike = style.font_strike

        # Color
        if style.font_color_rgb:
            font.color.rgb = style.font_color_rgb

    except Exception as e:
        logging.warning(f"Failed to apply style: {e}")


def _adjust_font_size_for_run(run: Run, style: Optional[RunStyle],
                              new_len: int, original_len: int):
    """Adjust run font size"""
    if not style or not style.font_size:
        return

    try:
        run.font.size = _calculate_adjusted_size(style.font_size, new_len, original_len)
    except Exception as e:
        logging.warning(f"Failed to adjust font size: {e}")


def _set_default_font(run: Run, target_lang: str):
    """Set default font"""
    try:
        if target_lang in ['中文', '日语', '韩语']:
            run.font.name = '微软雅黑'
            run._element.get_or_add_rPr()
            rFonts = run._element.rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            run.font.name = 'Times New Roman'
    except:
        pass


def _calculate_adjusted_size(original_size: Pt, new_len: int, original_len: int) -> Pt:
    """Calculate adjusted font size"""
    if not original_size:
        return Pt(11)

    if original_len == 0 or new_len <= original_len:
        return original_size

    ratio = new_len / original_len

    if ratio <= 1.2:
        return original_size
    elif ratio <= 1.5:
        return Pt(max(8, original_size.pt * 0.95))
    elif ratio <= 2.0:
        return Pt(max(8, original_size.pt * 0.90))
    else:
        return Pt(max(8, original_size.pt * 0.85))
